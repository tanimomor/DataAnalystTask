from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_formatted_doc(content):
    sections = [section.strip() for section in content.split('\n') if section.strip()]

    new_doc = Document()
    new_doc.add_paragraph("\n")
    for section in sections:
        p = new_doc.add_paragraph(f"============\n\n{section.strip()}\n")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return new_doc


def main():
    input_file = 'hadith.docx'
    output_file = 'hadith_formatted.docx'

    # Read the input .docx file
    doc = Document(input_file)

    content = ""
    for para in doc.paragraphs:
        content += para.text + "\n"

    formatted_doc = generate_formatted_doc(content)
    if formatted_doc is None:
        print(f"The {input_file} document is empty.")
    else:
        formatted_doc.save(output_file)
        print(f"Formatted content saved to {output_file}")


if __name__ == "__main__":
    main()
